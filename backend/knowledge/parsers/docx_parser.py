"""
DOCX parser — extracts text from .docx files using python-docx.

Splits documents by heading paragraphs (Heading 1/2/3 styles)
into logical sections. Falls back to full-text if no headings found.

Does NOT support old .doc format — convert to .docx first.
"""
import os
import logging


from .base import ParsedDocument, sanitize_id, detect_entity_from_filename, detect_category_from_path

logger = logging.getLogger("parser.docx")


def parse_docx(filepath: str) -> list[ParsedDocument]:
    """Parse a .docx file into a list of ParsedDocuments."""
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return []

    filename = os.path.basename(filepath)
    base_id = sanitize_id(filename)
    entity = detect_entity_from_filename(filename)
    category = detect_category_from_path(filepath)

    try:
        doc = Document(filepath)
    except Exception as e:
        logger.error("Failed to open DOCX '%s': %s", filepath, e)
        return []

    # Extract paragraphs with style info
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        is_heading = style_name.lower().startswith("heading")
        paragraphs.append((text, is_heading, style_name))

    if not paragraphs:
        logger.warning("DOCX '%s' has no text content", filepath)
        return []

    # Extract tables as markdown
    tables_text = []
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        if table_md:
            tables_text.append(table_md)

    # Split by headings
    sections = _split_by_headings(paragraphs)

    # If tables exist and weren't captured in sections, append them
    if tables_text and len(sections) <= 1:
        combined = sections[0][1] if sections else ""
        combined += "\n\n" + "\n\n".join(tables_text)
        sections = [(sections[0][0] if sections else "", combined)]

    documents = []
    if len(sections) <= 1:
        # No headings — single document
        full_text = sections[0][1] if sections else "\n".join(t for t, _, _ in paragraphs)
        documents.append(ParsedDocument(
            id=base_id,
            content=full_text,
            metadata={
                "source": "docx",
                "category": category,
                "topic": _infer_topic(full_text, filename),
                "doc_type": "docx",
                "entity": entity,
                "file_path": filepath,
            },
        ))
    else:
        for i, (title, content) in enumerate(sections):
            if len(content.strip()) < 30:
                continue
            documents.append(ParsedDocument(
                id=f"{base_id}_sec_{i}",
                content=f"## {title}\n\n{content}" if title else content,
                metadata={
                    "source": "docx",
                    "category": category,
                    "topic": _infer_topic(content, title or filename),
                    "doc_type": "docx",
                    "entity": entity,
                    "file_path": filepath,
                    "section_title": title or f"Section {i+1}",
                    "section_index": i,
                },
            ))

    logger.info("Parsed DOCX '%s': %d documents", filename, len(documents))
    return documents


def _split_by_headings(paragraphs: list[tuple[str, bool, str]]) -> list[tuple[str, str]]:
    """Split paragraphs into (title, content) sections by heading styles."""
    sections = []
    current_title = ""
    current_lines = []

    for text, is_heading, _ in paragraphs:
        if is_heading and current_lines:
            sections.append((current_title, "\n\n".join(current_lines)))
            current_title = text
            current_lines = []
        elif is_heading:
            current_title = text
        else:
            current_lines.append(text)

    if current_lines:
        sections.append((current_title, "\n\n".join(current_lines)))

    return sections


def _table_to_markdown(table) -> str:
    """Convert a docx table to markdown format."""
    rows = []
    for row in table.rows:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        rows.append(" | ".join(cells))

    if len(rows) < 2:
        return ""

    # Add separator after header
    header = rows[0]
    separator = " | ".join("---" for _ in rows[0].split(" | "))
    return "\n".join([header, separator] + rows[1:])


def _infer_topic(text: str, context: str = "") -> str:
    combined = (text[:500] + " " + context).lower()
    topic_map = {
        "financial_statements": ["balance sheet", "income", "cash flow", "资产负债", "利润表"],
        "valuation": ["p/e", "valuation", "dcf", "估值", "市盈率"],
        "earnings": ["earnings", "revenue", "quarterly", "财报", "营收"],
        "risk": ["risk", "volatility", "风险", "波动"],
        "macroeconomics": ["gdp", "inflation", "interest rate", "通胀", "利率"],
    }
    for topic, keywords in topic_map.items():
        if any(kw in combined for kw in keywords):
            return topic
    return "general"
